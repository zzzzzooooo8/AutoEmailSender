from __future__ import annotations

import logging
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from uuid import uuid4

from fastapi import UploadFile
from markitdown import MarkItDown

from app.core.config import get_settings

TEXT_EXTRACTABLE_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}
logger = logging.getLogger(__name__)


@dataclass(frozen=True, slots=True)
class StoredUpload:
    file_path: str
    original_name: str
    size_bytes: int
    sha256: str
    suffix: str


def build_display_name(file_name: str) -> str:
    candidate = Path(file_name).stem.strip()
    return candidate or "未命名材料"


def save_upload(file: UploadFile, *segments: str) -> StoredUpload:
    settings = get_settings()
    upload_dir = settings.uploads_dir.joinpath(*segments)
    upload_dir.mkdir(parents=True, exist_ok=True)

    original_name = Path(file.filename or "upload.bin").name
    suffix = Path(original_name).suffix.lower()
    target_path = upload_dir / f"{uuid4().hex}{suffix}"

    content = file.file.read()
    target_path.write_bytes(content)
    return StoredUpload(
        file_path=target_path.as_posix(),
        original_name=original_name,
        size_bytes=len(content),
        sha256=_hash_bytes(content),
        suffix=suffix,
    )


def delete_file(file_path: str | None) -> None:
    if not file_path:
        return
    target = Path(file_path)
    if target.exists():
        target.unlink()


def extract_text_from_document(file_path: str) -> str | None:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in TEXT_EXTRACTABLE_SUFFIXES:
        return None

    if suffix in {".txt", ".md"}:
        return _extract_text_file(path)

    content = _extract_text_with_markitdown(path)
    if content:
        return content

    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".pdf":
        return _extract_pdf_text(path)

    return None


def _extract_text_with_markitdown(path: Path) -> str | None:
    try:
        result = _get_markitdown().convert(path)
        content = (result.markdown or result.text_content or "").strip()
        return content or None
    except Exception:
        logger.exception("材料 Markdown 提取失败: %s", path.as_posix())
        return None


def _extract_text_file(path: Path) -> str | None:
    try:
        content = path.read_bytes()
    except OSError:
        logger.exception("材料文本读取失败: %s", path.as_posix())
        return None

    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = content.decode(encoding).strip()
            return text or None
        except UnicodeDecodeError:
            continue
    text = content.decode("utf-8", errors="ignore").strip()
    return text or None


def _extract_docx_text(path: Path) -> str | None:
    try:
        from docx import Document

        document = Document(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        content = "\n".join(item for item in [*paragraphs, *table_cells] if item).strip()
        return content or None
    except Exception:
        logger.exception("材料 DOCX 提取失败: %s", path.as_posix())
        return None


def _extract_pdf_text(path: Path) -> str | None:
    try:
        from pdfminer.high_level import extract_text

        content = (extract_text(path.as_posix()) or "").strip()
        if content:
            return content
    except Exception:
        logger.exception("材料 PDF pdfminer 提取失败: %s", path.as_posix())

    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        content = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
        return content or None
    except Exception:
        logger.exception("材料 PDF pypdf 提取失败: %s", path.as_posix())
        return None


def material_supports_text_extraction(file_name: str | None) -> bool:
    if not file_name:
        return False
    return Path(file_name).suffix.lower() in TEXT_EXTRACTABLE_SUFFIXES


def _hash_bytes(content: bytes) -> str:
    from hashlib import sha256

    return sha256(content).hexdigest()

@lru_cache(maxsize=1)
def _get_markitdown() -> MarkItDown:
    return MarkItDown(enable_plugins=False)
