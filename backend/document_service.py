#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换服务
处理docx文档转换为HTML格式，支持附件功能
"""

import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Tuple, Optional
import base64
from io import BytesIO
from backend.database import db
from backend.models.user_file import UserFile
import logging

logger = logging.getLogger(__name__)

class DocumentService:
    """文档转换服务类"""
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def docx_to_html(self, file_path: str) -> Dict[str, any]:
        """
        将docx文件转换为HTML格式
        
        Args:
            file_path: docx文件路径
            
        Returns:
            包含HTML内容和附件信息的字典
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            if not file_path.lower().endswith('.docx'):
                raise ValueError("仅支持.docx格式文件")
            
            # 读取docx文档
            doc = Document(file_path)
            
            # 转换为HTML
            html_content = self._convert_document_to_html(doc)
            
            # 提取图片和其他媒体文件
            attachments = self._extract_attachments(doc, file_path)
            
            return {
                'html_content': html_content,
                'attachments': attachments,
                'original_filename': os.path.basename(file_path)
            }
            
        except Exception as e:
            raise Exception(f"转换docx文件失败: {str(e)}")
    
    def _convert_document_to_html(self, doc: Document) -> str:
        """
        将Document对象转换为HTML字符串
        
        Args:
            doc: python-docx Document对象
            
        Returns:
            HTML字符串
        """
        html_parts = []
        
        # 按文档顺序处理所有元素
        # 使用document的element属性来获取正确的顺序
        try:
            # 遍历文档的所有元素，保持原始顺序
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    # 找到对应的段落对象
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            html_parts.append(self._convert_paragraph_to_html(paragraph))
                            break
                elif element.tag.endswith('tbl'):  # 表格
                    # 找到对应的表格对象
                    for table in doc.tables:
                        if table._element == element:
                            html_parts.append(self._convert_table_to_html(table))
                            break
        except Exception as e:
            # 如果无法按顺序处理，回退到原来的方法
            logger.warning(f"无法按文档顺序处理元素，使用备用方法: {e}")
            for paragraph in doc.paragraphs:
                html_parts.append(self._convert_paragraph_to_html(paragraph))
            
            for table in doc.tables:
                html_parts.append(self._convert_table_to_html(table))
        
        # 将内容包装在样式容器中，确保邮件发送时保持一致的格式（移除边框等视觉包装）
        content = '\n'.join(html_parts)
        wrapped_content = f'''<div style="font-family: 'Times New Roman', Times, serif; font-size: 12pt; line-height: 1.6; color: #333; word-wrap: break-word;">{content}</div>'''
        
        return wrapped_content
    
    def _convert_paragraph_to_html(self, paragraph) -> str:
        """
        将段落转换为HTML
        
        Args:
            paragraph: docx段落对象
            
        Returns:
            HTML字符串
        """
        if not paragraph.text.strip():
            return '<br>'
        
        # 处理段落样式
        style_attrs = []
        
        # 添加基本段落样式
        style_attrs.append('margin: 0')
        style_attrs.append('padding: 0')
        style_attrs.append('line-height: 1.5')
        
        # 对齐方式
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            style_attrs.append('text-align: center')
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            style_attrs.append('text-align: right')
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            style_attrs.append('text-align: justify')
        else:
            style_attrs.append('text-align: left')
        
        # 处理段落缩进
        if hasattr(paragraph.paragraph_format, 'first_line_indent') and paragraph.paragraph_format.first_line_indent:
            try:
                indent_pt = paragraph.paragraph_format.first_line_indent.pt
                style_attrs.append(f'text-indent: {indent_pt}pt')
            except:
                pass
        
        # 处理左缩进
        if hasattr(paragraph.paragraph_format, 'left_indent') and paragraph.paragraph_format.left_indent:
            try:
                left_indent_pt = paragraph.paragraph_format.left_indent.pt
                style_attrs.append(f'margin-left: {left_indent_pt}pt')
            except:
                pass
        
        # 处理右缩进
        if hasattr(paragraph.paragraph_format, 'right_indent') and paragraph.paragraph_format.right_indent:
            try:
                right_indent_pt = paragraph.paragraph_format.right_indent.pt
                style_attrs.append(f'margin-right: {right_indent_pt}pt')
            except:
                pass
        
        # 处理段落间距
        if hasattr(paragraph.paragraph_format, 'space_before') and paragraph.paragraph_format.space_before:
            try:
                space_before_pt = paragraph.paragraph_format.space_before.pt
                style_attrs.append(f'margin-top: {space_before_pt}pt')
            except:
                style_attrs.append('margin-top: 6pt')  # 默认段前间距
        else:
            style_attrs.append('margin-top: 6pt')
        
        if hasattr(paragraph.paragraph_format, 'space_after') and paragraph.paragraph_format.space_after:
            try:
                space_after_pt = paragraph.paragraph_format.space_after.pt
                style_attrs.append(f'margin-bottom: {space_after_pt}pt')
            except:
                style_attrs.append('margin-bottom: 6pt')  # 默认段后间距
        else:
            style_attrs.append('margin-bottom: 6pt')
        
        # 处理文本格式
        html_text = self._convert_runs_to_html(paragraph.runs)
        
        # 判断是否为标题
        if paragraph.style.name.startswith('Heading'):
            level = paragraph.style.name.replace('Heading ', '')
            try:
                level = int(level)
                if 1 <= level <= 6:
                    # 为标题添加额外样式
                    style_attrs.append('font-weight: bold')
                    if level == 1:
                        style_attrs.append('font-size: 18pt')
                    elif level == 2:
                        style_attrs.append('font-size: 16pt')
                    elif level == 3:
                        style_attrs.append('font-size: 14pt')
                    else:
                        style_attrs.append('font-size: 12pt')
                    
                    style_str = f' style="{";".join(style_attrs)}"' if style_attrs else ''
                    return f'<h{level}{style_str}>{html_text}</h{level}>'
            except ValueError:
                pass
        
        # 普通段落
        style_str = f' style="{";".join(style_attrs)}"' if style_attrs else ''
        return f'<p{style_str}>{html_text}</p>'
    
    def _convert_runs_to_html(self, runs) -> str:
        """
        将文本运行转换为HTML
        
        Args:
            runs: docx文本运行列表
            
        Returns:
            HTML字符串
        """
        html_parts = []
        
        for run in runs:
            text = run.text
            if not text:
                continue
            
            # 转义HTML特殊字符
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # 收集所有样式属性
            span_styles = []
            
            # 字体名称
            if run.font.name:
                span_styles.append(f'font-family: "{run.font.name}", serif')
            
            # 字体大小
            if run.font.size:
                try:
                    size_pt = run.font.size.pt
                    span_styles.append(f'font-size: {size_pt}pt')
                except:
                    pass
            
            # 字体颜色
            if run.font.color and run.font.color.rgb:
                try:
                    color = f'#{run.font.color.rgb}'
                    span_styles.append(f'color: {color}')
                except:
                    pass
            
            # 背景颜色
            if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                try:
                    bg_color = f'#{run.font.highlight_color}'
                    span_styles.append(f'background-color: {bg_color}')
                except:
                    pass
            
            # 应用样式 - 优先使用HTML语义标签，然后应用CSS样式
            # 首先应用HTML语义标签
            if run.bold:
                text = f'<strong>{text}</strong>'
            if run.italic:
                text = f'<em>{text}</em>'
            if run.underline:
                text = f'<u>{text}</u>'
            
            # 然后应用CSS样式（如果有的话）
            if span_styles:
                style_str = '; '.join(span_styles)
                text = f'<span style="{style_str}">{text}</span>'
            
            html_parts.append(text)
        
        return ''.join(html_parts)
    
    def _convert_table_to_html(self, table) -> str:
        """
        将表格转换为HTML
        
        Args:
            table: docx表格对象
            
        Returns:
            HTML字符串
        """
        # 改进的表格样式
        table_style = (
            'border-collapse: collapse; '
            'width: 100%; '
            'margin: 15px 0; '
            'font-family: inherit; '
            'border: 1px solid #000; '
            'background-color: #fff; '
            'font-size: 12pt;'
        )
        
        html_parts = [f'<table style="{table_style}">']
        
        for row_index, row in enumerate(table.rows):
            # 判断是否为表头行
            is_header = row_index == 0
            row_tag = 'th' if is_header else 'td'
            
            html_parts.append('<tr>')
            for cell in row.cells:
                cell_text = ''
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():  # 只处理非空段落
                        cell_text += self._convert_runs_to_html(paragraph.runs)
                        if paragraph != cell.paragraphs[-1]:  # 不是最后一个段落时添加换行
                            cell_text += '<br>'
                
                # 表头和普通单元格的不同样式
                if is_header:
                    cell_style = (
                        'border: 1px solid #000; '
                        'padding: 8px; '
                        'vertical-align: middle; '
                        'background-color: #f0f0f0; '
                        'font-weight: bold; '
                        'text-align: center; '
                        'font-size: 12pt;'
                    )
                else:
                    cell_style = (
                        'border: 1px solid #000; '
                        'padding: 8px; '
                        'vertical-align: top; '
                        'background-color: #fff; '
                        'font-size: 12pt;'
                    )
                
                html_parts.append(
                    f'<{row_tag} style="{cell_style}">{cell_text or "&nbsp;"}</{row_tag}>'
                )
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return '\n'.join(html_parts)
    
    def _extract_attachments(self, doc: Document, file_path: str) -> List[Dict[str, any]]:
        """
        提取文档中的附件（主要是图片）
        
        Args:
            doc: Document对象
            file_path: 原始文件路径
            
        Returns:
            附件信息列表
        """
        attachments = []
        
        # 添加原始docx文件作为附件
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
                file_base64 = base64.b64encode(file_content).decode('utf-8')
                
                attachments.append({
                    'filename': os.path.basename(file_path),
                    'content': file_base64,
                    'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'size': len(file_content)
                })
        except Exception as e:
            print(f"添加原始文件作为附件失败: {str(e)}")
        
        return attachments
    
    def validate_document(self, file_path: str) -> Dict[str, any]:
        """
        验证文档是否可以转换
        
        Args:
            file_path: 文件路径
            
        Returns:
            验证结果
        """
        result = {
            'valid': False,
            'message': '',
            'file_size': 0,
            'format': ''
        }
        
        try:
            if not os.path.exists(file_path):
                result['message'] = '文件不存在'
                return result
            
            file_size = os.path.getsize(file_path)
            result['file_size'] = file_size
            
            # 检查文件大小（限制为10MB）
            if file_size > 10 * 1024 * 1024:
                result['message'] = '文件大小超过10MB限制'
                return result
            
            # 检查文件格式
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.supported_formats:
                result['message'] = f'不支持的文件格式: {file_ext}'
                return result
            
            result['format'] = file_ext
            
            # 尝试打开文档
            doc = Document(file_path)
            result['valid'] = True
            result['message'] = '文档验证成功'
            
        except Exception as e:
            result['message'] = f'文档验证失败: {str(e)}'
        
        return result
    
    def get_document_preview(self, file_path: str, max_paragraphs: int = 3) -> str:
        """
        获取文档预览
        
        Args:
            file_path: 文件路径
            max_paragraphs: 最大段落数
            
        Returns:
            预览文本
        """
        try:
            doc = Document(file_path)
            preview_parts = []
            
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() and paragraph_count < max_paragraphs:
                    preview_parts.append(paragraph.text.strip())
                    paragraph_count += 1
            
            if paragraph_count == 0:
                return "文档内容为空或无法读取"
            
            preview = '\n\n'.join(preview_parts)
            if len(doc.paragraphs) > max_paragraphs:
                preview += '\n\n...（更多内容）'
            
            return preview
            
        except Exception as e:
            return f"预览失败: {str(e)}"
    
    def get_file_content(self, file_id: int, output_format: str = 'text') -> Tuple[Optional[str], Optional[str]]:
        """获取文件内容
        
        Args:
            file_id: 文件ID
            output_format: 输出格式 ('text' 或 'html')
        """
        try:
            # 获取文件记录
            user_file = UserFile.query.filter_by(id=file_id, is_active=True).first()
            if not user_file:
                return None, '文件不存在'
            
            # 检查文件是否存在
            if not os.path.exists(user_file.file_path):
                return None, '文件不存在于磁盘'
            
            # 根据文件类型处理内容
            file_ext = os.path.splitext(user_file.file_name)[1].lower()
            
            if file_ext in ['.docx']:
                # 处理docx文件
                try:
                    if output_format == 'html':
                        # 转换为HTML格式
                        result = self.docx_to_html(user_file.file_path)
                        return result['html_content'], None
                    else:
                        # 返回纯文本格式
                        doc = Document(user_file.file_path)
                        content_parts = []
                        
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                content_parts.append(paragraph.text.strip())
                        
                        content = '\n\n'.join(content_parts)
                        return content, None
                    
                except Exception as e:
                    logger.error(f'读取docx文件失败: {str(e)}')
                    return None, f'读取docx文件失败: {str(e)}'
            
            elif file_ext in ['.txt']:
                # 处理文本文件
                try:
                    with open(user_file.file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    if output_format == 'html':
                        # 将纯文本转换为HTML
                        content = content.replace('\n', '<br>\n')
                        content = f'<div>{content}</div>'
                    
                    return content, None
                except Exception as e:
                    logger.error(f'读取文本文件失败: {str(e)}')
                    return None, f'读取文本文件失败: {str(e)}'
            
            else:
                return None, f'不支持的文件类型: {file_ext}'
                
        except Exception as e:
            logger.error(f'获取文件内容失败: {str(e)}')
            return None, f'获取文件内容失败: {str(e)}'